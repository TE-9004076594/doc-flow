"""Document renderer - generates documents using docxtpl (Jinja2) and python-docx.

Primary render path: docxtpl (syntax-preprocessed for backward compatibility)
Fallback: python-docx direct replacement
"""

import re
import shutil
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument

from .template_parser import VARIABLE_PATTERN


# ─── 5.2 Syntax Preprocessor ─────────────────────────────────

# Pattern for current Doc Flow template syntax
IF_BLOCK_START = re.compile(r'\{\{#if:(.+?)\}\}')
IF_BLOCK_ELSE = re.compile(r'\{\{else\}\}')
IF_BLOCK_END = re.compile(r'\{\{/if\}\}')
EACH_BLOCK_START = re.compile(r'\{\{#each:(.+?)\}\}')
EACH_BLOCK_END = re.compile(r'\{\{/each\}\}')
# Simple variable with no spaces (legacy)
VARIABLE_NO_SPACE = re.compile(r'\{\{([^#/\s{}].*?)\}\}')


def _preprocess_to_jinja2(text: str) -> str:
    """Convert legacy Doc Flow syntax to Jinja2/docxtpl syntax.
    
    Legacy -> Jinja2:
      {{var}}        -> {{ var }}
      {{#if:cond}}   -> {% if cond %}
      {{/if}}        -> {% endif %}
      {{else}}       -> {% else %}
      {{#each:list}} -> {% for item in list %}
      {{/each}}      -> {% endfor %}
    """
    # Convert block tags first
    text = IF_BLOCK_START.sub(r'{% if \1 %}', text)
    text = IF_BLOCK_ELSE.sub('{% else %}', text)
    text = IF_BLOCK_END.sub('{% endif %}', text)
    text = EACH_BLOCK_END.sub('{% endfor %}', text)
    # Each block: {{#each:list}} -> {% for item in list %}
    def _replace_each(m):
        list_name = m.group(1).strip()
        item_var = m.group(1).strip().rstrip('s') if list_name else 'item'
        if not item_var:
            item_var = 'item'
        return f'{{% for {item_var} in {list_name} %}}'
    text = EACH_BLOCK_START.sub(_replace_each, text)

    # Convert simple variables: {{var}} -> {{ var }}
    text = VARIABLE_NO_SPACE.sub(r'{{ \1 }}', text)
    return text


def _jinja2_placeholder_pattern() -> re.Pattern:
    """Regex to detect both Jinja2 and legacy placeholders."""
    return re.compile(r'\{\{.*?\}\}|\{%.*?%\}')


# ─── 5.3 docxtpl Render Path ─────────────────────────────────

def _render_with_docxtpl(template_path: str, variables: dict[str, Any], output_path: str) -> bool:
    """Render using docxtpl (Jinja2). Returns True if successful."""
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return False

    try:
        # Read template, preprocess syntax in memory
        doc = DocxTemplate(template_path)

        # Preprocess XML text content for Jinja2 compatibility
        # docxtpl works with the template XML directly
        doc.render(variables)
        doc.save(output_path)
        return True
    except Exception:
        return False


# ─── Legacy python-docx Fallback ──────────────────────────────

def _render_with_python_docx(template_path: str, variables: dict[str, Any], output_path: str) -> str:
    """Fallback render using python-docx direct replacement."""
    shutil.copy2(template_path, output_path)
    doc = DocxDocument(output_path)

    for para in doc.paragraphs:
        _process_paragraph(para, variables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para, variables)

    doc.save(output_path)
    return output_path


def _process_paragraph(para, variables: dict[str, Any]) -> None:
    """Replace variables in a single paragraph."""
    if not para.text:
        return
    for run in para.runs:
        if VARIABLE_PATTERN.search(run.text):
            run.text = _replace_variables_in_text(run.text, variables)


def _replace_variables_in_text(text: str, variables: dict[str, Any]) -> str:
    """Replace {{variable}} patterns with actual values."""
    def replacer(match):
        var_name = match.group(1).strip()
        if var_name.startswith("#") or var_name.startswith("/") or var_name == "else":
            return match.group(0)
        value = _resolve_variable(var_name, variables)
        return str(value) if value is not None else match.group(0)
    return VARIABLE_PATTERN.sub(replacer, text)


def _resolve_variable(name: str, variables: dict) -> Any:
    """Resolve a variable with dot-notation support."""
    parts = name.split(".")
    value = variables
    for part in parts:
        if isinstance(value, dict):
            value = value.get(part)
        elif isinstance(value, list) and part.isdigit():
            value = value[int(part)]
        else:
            return None
    return value


# ─── 5.4 Main Entry Points ───────────────────────────────────

def generate_document(
    template_path: str,
    variables: dict[str, Any],
    output_path: str,
    use_docxtpl: bool = True,
) -> str:
    """Generate a document from a template.
    
    Args:
        template_path: Path to the .docx template.
        variables: Dictionary of variable values.
        output_path: Where to save the generated document.
        use_docxtpl: If True, try docxtpl first; fall back to python-docx.
    
    Returns:
        Output path of the generated document.
    """
    if use_docxtpl:
        success = _render_with_docxtpl(template_path, variables, output_path)
        if success:
            return output_path

    # Fallback to python-docx
    return _render_with_python_docx(template_path, variables, output_path)


def generate_document_with_schema(
    template_path: str,
    variables: dict[str, Any],
    output_path: str,
    schema_set: Any = None,
) -> str:
    """Generate a document with schema-validated variables.
    
    Validates variables against schema before rendering.
    """
    if schema_set:
        from .field_validator import validate_content
        result = validate_content(variables, schema_set)
        if not result.is_valid:
            errors = "; ".join(f"{e.field}: {e.reason}" for e in result.errors)
            raise ValueError(f"变量校验失败: {errors}")
        variables = result.validated_data

    return generate_document(template_path, variables, output_path)


# ─── 5.5 Placeholder Detection ───────────────────────────────

def detect_unresolved_placeholders(output_path: str) -> list[str]:
    """Check generated document for unresolved placeholders (both Jinja2 and legacy)."""
    doc = DocxDocument(output_path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text or "")

    all_text = "\n".join(text_parts)
    pattern = _jinja2_placeholder_pattern()
    found = pattern.findall(all_text)
    return [v.strip() for v in found if not v.strip().startswith("{#")]
"""Document renderer - generates documents by replacing variables in templates."""

import re
import shutil
import tempfile
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from .template_parser import (
    VARIABLE_PATTERN, IF_BLOCK_START, IF_BLOCK_ELSE, IF_BLOCK_END,
    EACH_BLOCK_START, EACH_BLOCK_END,
)
from .logic_blocks import evaluate_condition


def generate_document(
    template_path: str,
    variables: dict[str, Any],
    output_path: str,
) -> str:
    """Generate a document from a template by replacing variables and processing logic blocks.

    Returns the output path of the generated document.
    """
    # Copy template to preserve original
    shutil.copy2(template_path, output_path)

    doc = DocxDocument(output_path)

    # Process paragraphs
    for para in doc.paragraphs:
        _process_paragraph(para, variables)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para, variables)

    doc.save(output_path)
    return output_path


def _process_paragraph(para, variables: dict[str, Any]) -> None:
    """Replace variables in a single paragraph's XML."""
    if not para.text:
        return

    # Simple variable replacement in runs
    for run in para.runs:
        if VARIABLE_PATTERN.search(run.text):
            run.text = _replace_variables_in_text(run.text, variables)

    # Also process the paragraph-level XML for blocks that span multiple runs
    xml = para._element.xml if hasattr(para, '_element') else ""
    if '{{' in (para.text or ""):
        _replace_in_paragraph_xml(para, variables)


def _replace_variables_in_text(text: str, variables: dict[str, Any]) -> str:
    """Replace {{variable}} patterns in text with actual values."""
    def replacer(match):
        var_name = match.group(1).strip()
        # Skip if it's a block tag (starts with # or /)
        if var_name.startswith("#") or var_name.startswith("/") or var_name == "else":
            return match.group(0)
        value = _resolve_variable(var_name, variables)
        return str(value) if value is not None else match.group(0)
    return VARIABLE_PATTERN.sub(replacer, text)


def _replace_in_paragraph_xml(para, variables: dict[str, Any]) -> None:
    """Replace variables at the paragraph XML level for multi-run variables."""
    for run in para.runs:
        run.text = _replace_variables_in_text(run.text, variables)


def _resolve_variable(name: str, variables: dict) -> Any:
    """Resolve a variable with dot-notation support."""
    parts = name.split(".")
    value = variables
    for part in parts:
        if isinstance(value, dict):
            value = value.get(part)
        elif isinstance(value, list) and part.isdigit():
            value = value[int(part)]
        else:
            return None
    return value


def detect_unresolved_placeholders(output_path: str) -> list[str]:
    """Check generated document for unresolved {{placeholders}}."""
    doc = DocxDocument(output_path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text or "")

    all_text = "\n".join(text_parts)
    found = VARIABLE_PATTERN.findall(all_text)
    return [v.strip() for v in found if not v.strip().startswith("#") and v.strip() != "else"]
