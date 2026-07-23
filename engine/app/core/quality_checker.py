"""Quality checker - validates generated documents against template format specs."""

import re
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument

from .models import StyleSpec
from .extractor import extract_format
from .renderer import _jinja2_placeholder_pattern


# ─── 6.1 Style Consistency Check ─────────────────────────────

def check_page_setup(template_spec: StyleSpec, output_spec: StyleSpec) -> list[dict]:
    """Compare page setup between template and generated document."""
    issues: list[dict] = []
    t = template_spec.page
    o = output_spec.page

    checks = [
        ("页面尺寸", t.size, o.size),
        ("页面方向", t.orientation, o.orientation),
        ("上边距", f"{t.margin.top_cm:.2f}cm", f"{o.margin.top_cm:.2f}cm"),
        ("下边距", f"{t.margin.bottom_cm:.2f}cm", f"{o.margin.bottom_cm:.2f}cm"),
        ("左边距", f"{t.margin.left_cm:.2f}cm", f"{o.margin.left_cm:.2f}cm"),
        ("右边距", f"{t.margin.right_cm:.2f}cm", f"{o.margin.right_cm:.2f}cm"),
    ]
    for name, expected, actual in checks:
        if expected != actual:
            issues.append({
                "category": "page_setup",
                "item": name,
                "expected": expected,
                "actual": actual,
            })
    return issues


def check_styles(template_spec: StyleSpec, output_spec: StyleSpec) -> list[dict]:
    """Compare paragraph styles between template and generated document."""
    issues: list[dict] = []

    for style_name, t_style in template_spec.styles.items():
        o_style = output_spec.styles.get(style_name)
        if o_style is None:
            issues.append({
                "category": "style",
                "item": f"样式缺失: {style_name}",
                "expected": "存在",
                "actual": "缺失",
            })
            continue

        # Compare key style parameters
        param_checks = [
            ("字号", f"{t_style.size_pt}pt", f"{o_style.size_pt}pt"),
            ("粗体", str(t_style.bold), str(o_style.bold)),
            ("斜体", str(t_style.italic), str(o_style.italic)),
            ("对齐", t_style.alignment, o_style.alignment),
            ("行距", str(t_style.line_spacing), str(o_style.line_spacing)),
        ]
        for param, expected, actual in param_checks:
            if expected != actual:
                issues.append({
                    "category": "style",
                    "item": f"{style_name}.{param}",
                    "expected": expected,
                    "actual": actual,
                })

    return issues


def check_numbering(template_spec: StyleSpec, output_spec: StyleSpec) -> list[dict]:
    """Compare numbering definitions."""
    issues: list[dict] = []
    t_count = len(template_spec.numbering)
    o_count = len(output_spec.numbering)

    if t_count != o_count:
        issues.append({
            "category": "numbering",
            "item": "编号定义数量",
            "expected": str(t_count),
            "actual": str(o_count),
        })
    return issues


def check_table_format(template_spec: StyleSpec, output_spec: StyleSpec) -> list[dict]:
    """Compare table format."""
    issues: list[dict] = []
    t = template_spec.table
    o = output_spec.table

    if t.style_name != o.style_name:
        issues.append({
            "category": "table",
            "item": "表格样式",
            "expected": t.style_name,
            "actual": o.style_name,
        })
    return issues


# ─── 6.2 Placeholder Check ───────────────────────────────────

def check_unresolved_placeholders(output_path: str) -> list[dict]:
    """Check for unresolved placeholders in the generated document."""
    doc = DocxDocument(output_path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text or "")

    all_text = "\n".join(text_parts)
    pattern = _jinja2_placeholder_pattern()
    found = pattern.findall(all_text)
    if not found:
        return []
    return [{"category": "placeholder", "item": p.strip(), "expected": "已替换", "actual": "未替换"}
            for p in found if not p.strip().startswith("{#")]


# ─── 6.3 Quality Report ──────────────────────────────────────

def calculate_consistency_score(total_checks: int, passed_checks: int) -> float:
    """Calculate format consistency percentage."""
    if total_checks == 0:
        return 100.0
    return round((passed_checks / total_checks) * 100, 2)


class QualityReport:
    """Quality check result for a generated document."""

    def __init__(self):
        self.page_setup_issues: list[dict] = []
        self.style_issues: list[dict] = []
        self.numbering_issues: list[dict] = []
        self.table_issues: list[dict] = []
        self.placeholder_issues: list[dict] = []
        self.consistency_score: float = 100.0

    @property
    def total_issues(self) -> int:
        return (len(self.page_setup_issues) + len(self.style_issues)
                + len(self.numbering_issues) + len(self.table_issues)
                + len(self.placeholder_issues))

    @property
    def passed(self) -> bool:
        return self.total_issues == 0

    def to_dict(self) -> dict:
        return {
            "passed": self.passed,
            "consistency_score": self.consistency_score,
            "total_issues": self.total_issues,
            "details": {
                "page_setup": {"passed": len(self.page_setup_issues) == 0, "issues": self.page_setup_issues},
                "styles": {"passed": len(self.style_issues) == 0, "issues": self.style_issues},
                "numbering": {"passed": len(self.numbering_issues) == 0, "issues": self.numbering_issues},
                "table": {"passed": len(self.table_issues) == 0, "issues": self.table_issues},
                "placeholders": {"passed": len(self.placeholder_issues) == 0, "issues": self.placeholder_issues},
            },
        }


# ─── Main Entry ──────────────────────────────────────────────

def run_quality_check(template_path: str, output_path: str, threshold: float = 98.0) -> QualityReport:
    """Run full quality check on a generated document against its template."""
    report = QualityReport()

    try:
        template_spec = extract_format(template_path)
        output_spec = extract_format(output_path)
    except Exception as e:
        report.page_setup_issues.append({
            "category": "error", "item": "格式提取失败", "expected": "", "actual": str(e),
        })
        return report

    # Run all checks
    report.page_setup_issues = check_page_setup(template_spec, output_spec)
    report.style_issues = check_styles(template_spec, output_spec)
    report.numbering_issues = check_numbering(template_spec, output_spec)
    report.table_issues = check_table_format(template_spec, output_spec)
    report.placeholder_issues = check_unresolved_placeholders(output_path)

    # Calculate score
    all_issues = (report.page_setup_issues + report.style_issues
                  + report.numbering_issues + report.table_issues
                  + report.placeholder_issues)
    total_check_points = 50  # arbitrary baseline for scoring
    passed = total_check_points - len(all_issues)
    report.consistency_score = calculate_consistency_score(total_check_points, max(0, passed))
    report.consistency_score = min(100.0, report.consistency_score)

    return report
