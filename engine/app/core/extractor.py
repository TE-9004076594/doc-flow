"""Word format extractor - extracts complete formatting specs from .docx templates.

Extracts page setup, styles, numbering, tables, headers/footers using
python-docx (high-level) and lxml (deep OOXML access).
"""

import hashlib
import json
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches, Emu
from lxml import etree

from .models import (
    StyleSpec,
    PageSetup,
    PageMargin,
    ParagraphStyle,
    NumberingDefinition,
    NumberingLevel,
    TableFormat,
    HeaderFooter,
)


def extract_format(template_path: str) -> StyleSpec:
    """Extract complete format specification from a .docx template."""
    doc = DocxDocument(template_path)
    spec = StyleSpec()

    spec.page = _extract_page_setup(doc)
    spec.styles = _extract_styles(doc)
    spec.numbering = _extract_numbering(doc)
    spec.table = _extract_table_format(doc)
    header, footer = _extract_header_footer(doc)
    spec.header = header
    spec.footer = footer
    spec.constraints = _generate_constraints(spec)

    return spec


# ─── 2.1 Page Setup ──────────────────────────────────────────

def _extract_page_setup(doc: DocxDocument) -> PageSetup:
    """Extract page-level settings from document sections."""
    setup = PageSetup()

    for section in doc.sections:
        # Paper size
        page_w = section.page_width
        page_h = section.page_height
        if page_w and page_h:
            w_cm = _emu_to_cm(page_w)
            h_cm = _emu_to_cm(page_h)
            setup.size = _detect_paper_size(w_cm, h_cm)
            setup.orientation = "landscape" if w_cm > h_cm else "portrait"

        # Margins
        setup.margin = PageMargin(
            top_cm=_emu_to_cm(section.top_margin) if section.top_margin else 2.54,
            bottom_cm=_emu_to_cm(section.bottom_margin) if section.bottom_margin else 2.54,
            left_cm=_emu_to_cm(section.left_margin) if section.left_margin else 3.17,
            right_cm=_emu_to_cm(section.right_margin) if section.right_margin else 3.17,
        )

        # Header / footer distance
        if section.header_distance is not None:
            setup.header_distance_cm = _emu_to_cm(section.header_distance)
        if section.footer_distance is not None:
            setup.footer_distance_cm = _emu_to_cm(section.footer_distance)

        # Columns via lxml (python-docx doesn't expose this)
        sect_pr = section._sectPr
        if sect_pr is not None:
            cols_elem = sect_pr.find(qn('w:cols'))
            if cols_elem is not None:
                num_cols = cols_elem.get(qn('w:num'))
                if num_cols:
                    try:
                        setup.columns = int(num_cols)
                    except (ValueError, TypeError):
                        pass

        # Only extract from first section (primary) for MVP
        break

    return setup


def _emu_to_cm(emu: Any) -> float:
    """Convert EMU (English Metric Unit) to centimeters."""
    if emu is None:
        return 0.0
    return float(emu) / 914400 * 2.54


def _detect_paper_size(w_cm: float, h_cm: float) -> str:
    """Detect standard paper size from dimensions in cm."""
    sizes = {
        (21.0, 29.7): "A4",
        (29.7, 21.0): "A4",
        (14.8, 21.0): "A5",
        (21.0, 14.8): "A5",
        (42.0, 29.7): "A3",
        (29.7, 42.0): "A3",
        (21.6, 27.9): "Letter",
        (27.9, 21.6): "Letter",
        (21.6, 35.6): "Legal",
        (35.6, 21.6): "Legal",
    }
    key = (round(w_cm, 1), round(h_cm, 1))
    return sizes.get(key, f"{w_cm:.1f}x{h_cm:.1f}cm")


# ─── 2.2 Style Extraction ─────────────────────────────────────

def _extract_styles(doc: DocxDocument) -> dict[str, ParagraphStyle]:
    """Extract paragraph and character style definitions."""
    styles: dict[str, ParagraphStyle] = {}

    for style in doc.styles:
        if style.type is None:
            continue
        # We primarily care about paragraph styles
        from docx.enum.style import WD_STYLE_TYPE
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue

        ps = ParagraphStyle()
        ps.bold = style.font.bold or False
        ps.italic = style.font.italic or False

        # Font names (Chinese / English)
        if style.font.name:
            ps.font_en = style.font.name
        # Try to get Chinese font via style element
        rpr = style.element.find(qn('w:rPr'))
        if rpr is not None:
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is not None:
                east_asia = rfonts.get(qn('w:eastAsia'))
                if east_asia:
                    ps.font_cn = east_asia
                if not ps.font_en:
                    ps.font_en = rfonts.get(qn('w:ascii')) or rfonts.get(qn('w:hAnsi')) or ""

        # Font size
        if style.font.size:
            ps.size_pt = style.font.size.pt
        elif rpr is not None:
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                if val:
                    try:
                        ps.size_pt = int(val) / 2  # half-points
                    except (ValueError, TypeError):
                        pass

        # Font color
        if style.font.color and style.font.color.rgb:
            ps.color_hex = str(style.font.color.rgb)

        # Paragraph format
        pf = style.paragraph_format
        if pf.alignment is not None:
            align_map = {
                0: "left", 1: "center", 2: "right", 3: "justify",
                4: "start", 5: "end", 6: "distribute",
            }
            ps.alignment = align_map.get(pf.alignment, "left")

        if pf.first_line_indent:
            ps.first_line_indent_chars = _indent_to_chars(pf.first_line_indent, ps.size_pt)

        if pf.line_spacing:
            ps.line_spacing = pf.line_spacing

        if pf.space_before:
            ps.space_before_pt = pf.space_before.pt
        if pf.space_after:
            ps.space_after_pt = pf.space_after.pt

        ps.keep_with_next = pf.keep_with_next or False
        # keep_lines not available in all python-docx versions
        ps.page_break_before = pf.page_break_before or False

        styles[style.name] = ps

    return styles


def _indent_to_chars(indent: Any, font_size_pt: float) -> int:
    """Convert indentation to approximate character count."""
    if indent is None or font_size_pt <= 0:
        return 0
    try:
        indent_cm = indent.cm if hasattr(indent, 'cm') else float(indent) / 914400 * 2.54
        char_width_cm = font_size_pt * 0.035  # rough estimate
        if char_width_cm > 0:
            return max(0, int(round(indent_cm / char_width_cm)))
    except (ValueError, TypeError, AttributeError):
        pass
    return 0


# ─── 2.3 Numbering Extraction ─────────────────────────────────

def _extract_numbering(doc: DocxDocument) -> list[NumberingDefinition]:
    """Extract multi-level numbering definitions via lxml (deep OOXML)."""
    definitions: list[NumberingDefinition] = []

    # Access the numbering part from the document's package
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return definitions

    if numbering_part is None:
        return definitions

    numbering_elem = numbering_part._element
    if numbering_elem is None:
        return definitions

    # Parse abstract numbering definitions
    for abs_num in numbering_elem.findall(qn('w:abstractNum')):
        abs_id_str = abs_num.get(qn('w:abstractNumId'))
        if abs_id_str is None:
            continue
        try:
            abs_id = int(abs_id_str)
        except ValueError:
            continue

        nd = NumberingDefinition(abstract_num_id=abs_id)

        for lvl in abs_num.findall(qn('w:lvl')):
            lvl_num_str = lvl.get(qn('w:ilvl'))
            try:
                lvl_num = int(lvl_num_str) if lvl_num_str else 0
            except ValueError:
                continue

            nl = NumberingLevel(level=lvl_num)

            # Number format
            num_fmt = lvl.find(qn('w:numFmt'))
            if num_fmt is not None:
                nl.number_format = num_fmt.get(qn('w:val')) or "decimal"

            # Number template (e.g. "%1.%2")
            lvl_text = lvl.find(qn('w:lvlText'))
            if lvl_text is not None:
                nl.template = lvl_text.get(qn('w:val')) or f"%{lvl_num + 1}."

            # Start value
            start = lvl.find(qn('w:start'))
            if start is not None:
                start_val = start.get(qn('w:val'))
                if start_val:
                    try:
                        nl.start = int(start_val)
                    except ValueError:
                        pass

            # Indentation
            ppr = lvl.find(qn('w:pPr'))
            if ppr is not None:
                ind = ppr.find(qn('w:ind'))
                if ind is not None:
                    left = ind.get(qn('w:left'))
                    if left:
                        try:
                            nl.indent_left_cm = float(left) / 567  # twips to cm
                        except ValueError:
                            pass
                    hang = ind.get(qn('w:hanging'))
                    if hang:
                        try:
                            nl.indent_hanging_cm = float(hang) / 567
                        except ValueError:
                            pass

            # Font
            rpr = lvl.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    nl.font = rfonts.get(qn('w:ascii')) or rfonts.get(qn('w:hAnsi')) or ""
                sz = rpr.find(qn('w:sz'))
                if sz is not None:
                    val = sz.get(qn('w:val'))
                    if val:
                        try:
                            nl.size_pt = int(val) / 2
                        except ValueError:
                            pass

            nd.levels.append(nl)

        definitions.append(nd)

    return definitions


# ─── 2.4 Table Format Extraction ──────────────────────────────

def _extract_table_format(doc: DocxDocument) -> TableFormat:
    """Extract table formatting from the first table found."""
    tf = TableFormat()

    if not doc.tables:
        return tf

    table = doc.tables[0]

    # Table style
    if table.style and table.style.name:
        tf.style_name = table.style.name

    # Try to get detailed table properties via lxml
    tbl_elem = table._tbl
    tbl_pr = tbl_elem.find(qn('w:tblPr'))
    if tbl_pr is not None:
        # Borders
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is not None:
            for border_name in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
                border = tbl_borders.find(qn(f'w:{border_name}'))
                if border is not None:
                    sz = border.get(qn('w:sz'))
                    if sz:
                        try:
                            tf.border_size_pt = int(sz) / 8  # 1/8 pt units
                        except ValueError:
                            pass
                    color = border.get(qn('w:color'))
                    if color:
                        tf.border_color_hex = color

        # Cell margins
        tbl_cell_mar = tbl_pr.find(qn('w:tblCellMar'))
        if tbl_cell_mar is not None:
            top = tbl_cell_mar.find(qn('w:top'))
            if top is not None:
                w_val = top.get(qn('w:w'))
                if w_val:
                    try:
                        tf.cell_padding_pt = int(w_val) / 20  # twips to pt
                    except ValueError:
                        pass

    # Check first row for header formatting
    if table.rows:
        first_row = table.rows[0]
        for cell in first_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.bold:
                        tf.header_bold = True
                        break

        # Shading on first row cells
        for cell in first_row.cells:
            tc_pr = cell._tc.find(qn('w:tcPr'))
            if tc_pr is not None:
                shade = tc_pr.find(qn('w:shd'))
                if shade is not None:
                    fill = shade.get(qn('w:fill'))
                    if fill and fill != "auto":
                        tf.shading_color_hex = fill
                        break

    return tf


# ─── 2.5 Header / Footer Extraction ───────────────────────────

def _extract_header_footer(doc: DocxDocument) -> tuple[HeaderFooter, HeaderFooter]:
    """Extract header and footer configuration."""
    header = HeaderFooter()
    footer = HeaderFooter()

    for section in doc.sections:
        sect_pr = section._sectPr

        # Different first page
        title_pg = sect_pr.find(qn('w:titlePg'))
        if title_pg is not None:
            header.different_first_page = True
            footer.different_first_page = True

        # Odd/even differentiation (via lxml)
        even_or_odd = False
        for pg_ref_type in ('headerReference', 'footerReference'):
            for ref in sect_pr.findall(qn(f'w:{pg_ref_type}')):
                if ref is not None:
                    hdr_type = ref.get(qn('w:type'))
                    if hdr_type in ('even', 'default'):
                        even_or_odd = True

        header.different_odd_even = even_or_odd
        footer.different_odd_even = even_or_odd

        # Extract header content
        try:
            if section.header:
                header.content = section.header.paragraphs[0].text if section.header.paragraphs else ""
            if section.even_page_header:
                pass  # content stored separately
        except Exception:
            pass

        # Extract footer content
        try:
            if section.footer:
                footer.content = section.footer.paragraphs[0].text if section.footer.paragraphs else ""
        except Exception:
            pass

        # First section only
        break

    return header, footer


# ─── 2.6 Fingerprint ──────────────────────────────────────────

def compute_fingerprint(spec: StyleSpec) -> str:
    """Compute a SHA-256 fingerprint of the format specification."""
    return spec.fingerprint()


# ─── Constraints ──────────────────────────────────────────────

def _generate_constraints(spec: StyleSpec) -> list[str]:
    """Generate format constraints from extracted spec."""
    constraints = [
        "禁止新增样式",
        "禁止修改既有编号层级",
        "仅允许替换占位符",
        "禁止修改页面设置",
        "禁止修改表格样式",
    ]
    if spec.header.content:
        constraints.append("保留页眉内容")
    if spec.footer.content:
        constraints.append("保留页脚内容")
    return constraints
