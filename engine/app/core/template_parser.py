"""Template parser - extracts variables and block structures from .docx OOXML."""

import re
from typing import Optional
from xml.etree import ElementTree as ET

from docx import Document as DocxDocument
from docx.oxml.ns import qn


# Patterns for template syntax
VARIABLE_PATTERN = re.compile(r'\{\{([^#/].*?)\}\}')
IF_BLOCK_START = re.compile(r'\{\{#if:(.+?)\}\}')
IF_BLOCK_ELSE = re.compile(r'\{\{else\}\}')
IF_BLOCK_END = re.compile(r'\{\{/if\}\}')
EACH_BLOCK_START = re.compile(r'\{\{#each:(.+?)\}\}')
EACH_BLOCK_END = re.compile(r'\{\{/each\}\}')


class ParsedVariable:
    """Represents a detected template variable."""

    def __init__(self, name: str, context: Optional[dict] = None):
        self.name = name.strip()
        self.context = context or {}


class ParsedTemplate:
    """Represents a fully parsed template with variables and blocks."""

    def __init__(self):
        self.variables: list[ParsedVariable] = []
        self.has_condition_blocks: bool = False
        self.has_loop_blocks: bool = False
        self.block_count: int = 0


def parse_template(file_path: str) -> ParsedTemplate:
    """Parse a .docx template file to extract variables and structure metadata."""
    doc = DocxDocument(file_path)
    result = ParsedTemplate()

    seen_vars: set[str] = set()
    all_text = _extract_all_text(doc)

    # Extract simple variables
    for match in VARIABLE_PATTERN.finditer(all_text):
        var_name = match.group(1).strip()
        if var_name and var_name not in seen_vars:
            seen_vars.add(var_name)
            result.variables.append(ParsedVariable(var_name))

    # Detect logic blocks
    if IF_BLOCK_START.search(all_text):
        result.has_condition_blocks = True
        result.block_count += len(IF_BLOCK_START.findall(all_text))

    if EACH_BLOCK_START.search(all_text):
        result.has_loop_blocks = True
        result.block_count += len(EACH_BLOCK_START.findall(all_text))

    return result


def extract_variables_from_paragraph(paragraph) -> list[str]:
    """Extract variable names from a single paragraph."""
    text = _get_paragraph_text(paragraph)
    return [m.group(1).strip() for m in VARIABLE_PATTERN.finditer(text)]


def _extract_all_text(doc: DocxDocument) -> str:
    """Extract all text from a document for pattern matching."""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text or "")
    return "\n".join(texts)


def _get_paragraph_text(paragraph) -> str:
    """Get full text from a paragraph including runs."""
    return paragraph.text or ""


def validate_template_structure(file_path: str) -> list[str]:
    """Validate template block syntax. Returns list of errors."""
    doc = DocxDocument(file_path)
    all_text = _extract_all_text(doc)
    errors: list[str] = []

    # Check balanced if/endif blocks
    if_starts = len(IF_BLOCK_START.findall(all_text))
    if_ends = len(IF_BLOCK_END.findall(all_text))
    if if_starts != if_ends:
        errors.append(f"条件块不匹配: {if_starts} 个 {{#if}} 但 {if_ends} 个 {{/if}}")

    # Check balanced each/end blocks
    each_starts = len(EACH_BLOCK_START.findall(all_text))
    each_ends = len(EACH_BLOCK_END.findall(all_text))
    if each_starts != each_ends:
        errors.append(f"循环块不匹配: {each_starts} 个 {{#each}} 但 {each_ends} 个 {{/each}}")

    return errors
