"""Tests for format extractor and data models."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.core.models import (
    StyleSpec, PageSetup, PageMargin, ParagraphStyle,
    TableFormat, FieldSchema, FieldSchemaSet, FieldType,
)
from app.core.extractor import extract_format, compute_fingerprint


def _create_test_template() -> str:
    """Create a .docx file with diverse formatting for testing."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Custom styles
    style = doc.styles.add_style('TestHeading', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = None  # will use default
    from docx.shared import RGBColor
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5

    # Content with placeholder
    para = doc.add_paragraph("合同编号：{{contract_id}}", style='Normal')
    para = doc.add_paragraph("甲方：{{customer_name}}", style='TestHeading')

    # Table
    table = doc.add_table(rows=2, cols=3, style='Table Grid')
    table.cell(0, 0).text = '项目'
    table.cell(0, 1).text = '金额'
    table.cell(0, 2).text = '{{amount}}'

    doc.save(tmp.name)
    return tmp.name


class TestStyleSpec:
    """Test StyleSpec data model."""

    def test_fingerprint_generation(self):
        spec = StyleSpec(template_id="test-001")
        fp = spec.fingerprint()
        assert isinstance(fp, str)
        assert len(fp) == 16  # SHA-256 truncated to 16 chars

    def test_fingerprint_changes_with_content(self):
        spec1 = StyleSpec(template_id="test-001")
        spec2 = StyleSpec(template_id="test-002")
        assert spec1.fingerprint() != spec2.fingerprint()


class TestFieldSchemaSet:
    """Test FieldSchemaSet JSON Schema generation."""

    def test_to_json_schema(self):
        schema_set = FieldSchemaSet(template_id="test-001")
        schema_set.fields = [
            FieldSchema(name="customer_name", field_type=FieldType.STRING, required=True, max_length=100),
            FieldSchema(name="amount", field_type=FieldType.NUMBER, required=True, min_value=0),
            FieldSchema(name="department", field_type=FieldType.ENUM, enum_values=["HR", "Sales"]),
        ]
        json_schema = schema_set.to_json_schema()
        assert json_schema["type"] == "object"
        assert "customer_name" in json_schema["properties"]
        assert json_schema["properties"]["customer_name"]["maxLength"] == 100
        assert json_schema["properties"]["amount"]["minimum"] == 0
        assert json_schema["properties"]["department"]["enum"] == ["HR", "Sales"]
        assert "customer_name" in json_schema["required"]
        assert "amount" in json_schema["required"]

    def test_optional_field_not_in_required(self):
        schema_set = FieldSchemaSet(template_id="test-001")
        schema_set.fields = [
            FieldSchema(name="required_field", required=True),
            FieldSchema(name="optional_field", required=False),
        ]
        json_schema = schema_set.to_json_schema()
        assert "required_field" in json_schema["required"]
        assert "optional_field" not in json_schema["required"]


class TestExtractor:
    """Test format extraction from .docx files."""

    def setup_method(self):
        self.template_path = _create_test_template()

    def teardown_method(self):
        Path(self.template_path).unlink(missing_ok=True)

    def test_extract_page_setup(self):
        spec = extract_format(self.template_path)
        assert spec.page.size == "A4"
        assert spec.page.orientation == "portrait"
        assert abs(spec.page.margin.top_cm - 2.54) < 0.1
        assert abs(spec.page.margin.left_cm - 3.17) < 0.1

    def test_extract_styles(self):
        spec = extract_format(self.template_path)
        assert "TestHeading" in spec.styles
        hs = spec.styles["TestHeading"]
        assert hs.bold is True
        assert hs.size_pt == 16.0
        assert hs.line_spacing == 1.5

    def test_extract_normal_style(self):
        spec = extract_format(self.template_path)
        assert "Normal" in spec.styles
        ns = spec.styles["Normal"]
        assert ns.size_pt == 12.0

    def test_extract_table_format(self):
        spec = extract_format(self.template_path)
        assert spec.table.style_name == "Table Grid"
        # At minimum detect table presence
        assert spec.table is not None

    def test_fingerprint_consistency(self):
        spec1 = extract_format(self.template_path)
        spec2 = extract_format(self.template_path)
        assert spec1.fingerprint() == spec2.fingerprint()

    def test_extract_constraints(self):
        spec = extract_format(self.template_path)
        assert len(spec.constraints) > 0
        assert "禁止新增样式" in spec.constraints
