"""Format fidelity tests - verify generated document preserves template formatting."""

import os
import tempfile
import pytest
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.renderer import generate_document


def _create_styled_template() -> str:
    """Create a .docx template with known styling for fidelity testing."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()

    # Add a heading
    heading = doc.add_heading("测试标题：{{title}}", level=1)

    # Add a formatted paragraph
    p = doc.add_paragraph()
    run = p.add_run("公司名称：{{company_name}}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Add a normal paragraph
    p2 = doc.add_paragraph("日期：{{sign_date}}")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add a list item
    p3 = doc.add_paragraph("金额：{{amount}} 元")
    p3.paragraph_format.first_line_indent = Inches(0.3)

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.text = "{{cell_content}}"

    doc.save(tmp.name)
    return tmp.name


class TestFormatFidelity:
    """Verify generated documents preserve template formatting."""

    def test_heading_style_preserved(self):
        """Heading style should be preserved after variable replacement."""
        template_path = _create_styled_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "title": "年度报告",
            "company_name": "Acme Corp",
            "sign_date": "2024-01-01",
            "amount": "1,000,000",
            "cell_content": "数据",
        }, output_path)

        gen_doc = DocxDocument(output_path)

        # Check heading style preserved
        heading = gen_doc.paragraphs[0]
        assert "年度报告" in heading.text, "Heading variable should be replaced"
        assert "Heading" in (heading.style.name or ""), "Heading style should be preserved"

        # Cleanup
        os.unlink(template_path)
        os.unlink(output_path)

    def test_bold_text_preserved(self):
        """Bold formatting should be preserved after variable replacement."""
        template_path = _create_styled_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "title": "报告",
            "company_name": "Acme Corp",
            "sign_date": "2024-01-01",
            "amount": "500,000",
            "cell_content": "x",
        }, output_path)

        gen_doc = DocxDocument(output_path)
        second_para = gen_doc.paragraphs[1]

        # Check bold runs exist and contain replaced text
        bold_found = any(run.bold for run in second_para.runs)
        assert bold_found, "Bold formatting should be preserved"
        assert "Acme Corp" in second_para.text, "Variable should be replaced"

        os.unlink(template_path)
        os.unlink(output_path)

    def test_alignment_preserved(self):
        """Paragraph alignment should be preserved."""
        template_path = _create_styled_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "title": "报告",
            "company_name": "Acme",
            "sign_date": "2024-06-15",
            "amount": "100",
            "cell_content": "d",
        }, output_path)

        gen_doc = DocxDocument(output_path)

        # The right-aligned paragraph should still be right-aligned
        right_aligned = None
        for p in gen_doc.paragraphs:
            if "2024-06-15" in (p.text or ""):
                right_aligned = p
                break

        assert right_aligned is not None, "Date paragraph should exist"
        assert right_aligned.alignment == 2, "Right alignment (2) should be preserved"

        os.unlink(template_path)
        os.unlink(output_path)

    def test_table_structure_preserved(self):
        """Table structure (rows, cols) should be preserved."""
        template_path = _create_styled_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "title": "t",
            "company_name": "c",
            "sign_date": "d",
            "amount": "a",
            "cell_content": "表格数据",
        }, output_path)

        gen_doc = DocxDocument(output_path)

        assert len(gen_doc.tables) > 0, "Table should be preserved"
        table = gen_doc.tables[0]
        assert len(table.rows) == 2, "Table should have 2 rows"
        assert len(table.columns) == 2, "Table should have 2 columns"

        os.unlink(template_path)
        os.unlink(output_path)

    def test_font_size_preserved(self):
        """Font size should be preserved after variable replacement."""
        template_path = _create_styled_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "title": "报告",
            "company_name": "Acme Corp",
            "sign_date": "2024-01-01",
            "amount": "500,000",
            "cell_content": "x",
        }, output_path)

        gen_doc = DocxDocument(output_path)
        second_para = gen_doc.paragraphs[1]

        for run in second_para.runs:
            if "Acme Corp" in (run.text or ""):
                assert run.font.size is not None, "Font size should be preserved"

        os.unlink(template_path)
        os.unlink(output_path)

    def test_no_format_degradation(self):
        """Generated document should not introduce extra formatting artifacts."""
        template_path = _create_styled_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "title": "报告",
            "company_name": "Acme Corp",
            "sign_date": "2024-12-31",
            "amount": "999",
            "cell_content": "v",
        }, output_path)

        gen_doc = DocxDocument(output_path)
        raw_doc = DocxDocument(template_path)

        # Should have same number of paragraphs (variable replacements don't add/remove paragraphs)
        assert len(gen_doc.paragraphs) == len(raw_doc.paragraphs), (
            f"Paragraph count should match: {len(gen_doc.paragraphs)} vs {len(raw_doc.paragraphs)}"
        )

        os.unlink(template_path)
        os.unlink(output_path)
