"""Performance tests - verify generation time and concurrent processing."""

import os
import tempfile
import time
import threading
import pytest
from docx import Document as DocxDocument

from app.core.renderer import generate_document


def _create_template() -> str:
    """Create a template with multiple variables for performance testing."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()

    # Create a substantial template (10 paragraphs with variables)
    for i in range(10):
        p = doc.add_paragraph()
        p.add_run(f"段落{i+1}：{{var{i}}} 内容")

    # Add a table with variables
    table = doc.add_table(rows=5, cols=3)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "{{cell_data}}"

    doc.save(tmp.name)
    return tmp.name


def _create_large_template(num_paragraphs: int = 50, num_vars: int = 20) -> str:
    """Create a large template for stress testing."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()

    for i in range(num_paragraphs):
        doc.add_paragraph(f"第{i+1}段：{{var{i % num_vars}}} 值{{value_{i % num_vars}}}")

    doc.save(tmp.name)
    return tmp.name


class TestGenerationPerformance:
    """Verify document generation meets performance targets."""

    def test_single_generation_within_timeout(self):
        """Single document generation should complete within 10 seconds."""
        template_path = _create_template()
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        variables = {f"var{i}": f"测试值{i}" for i in range(10)}

        start = time.time()
        generate_document(template_path, variables, output_path)
        elapsed = time.time() - start

        assert elapsed < 10.0, (
            f"Single generation took {elapsed:.2f}s, expected < 10s"
        )
        print(f"  ✓ Single generation: {elapsed:.3f}s (limit: 10s)")

        os.unlink(template_path)
        os.unlink(output_path)

    def test_large_template_generation(self):
        """Large template (50 paragraphs) should generate within time limit."""
        template_path = _create_large_template(50, 20)
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        variables = {f"var{i}": f"V{i}" for i in range(20)}
        variables.update({f"value_{i}": f"数据点{i}" for i in range(20)})

        start = time.time()
        generate_document(template_path, variables, output_path)
        elapsed = time.time() - start

        assert elapsed < 15.0, (
            f"Large template generation took {elapsed:.2f}s, expected < 15s"
        )
        print(f"  ✓ Large template (50 paragraphs): {elapsed:.3f}s")

        os.unlink(template_path)
        os.unlink(output_path)

    def test_batch_concurrent_generation(self):
        """Multiple concurrent generations should complete correctly."""
        import concurrent.futures

        template_path = _create_template()
        variables_set = [
            {f"var{i}": f"批次{j}_值{i}" for i in range(10)}
            for j in range(5)
        ]

        outputs = []
        for _ in range(5):
            out = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            outputs.append(out.name)
            out.close()

        start = time.time()
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            futures = [
                executor.submit(generate_document, template_path, vars, out_path)
                for vars, out_path in zip(variables_set, outputs)
            ]
            for future in concurrent.futures.as_completed(futures):
                future.result()

        total_time = time.time() - start

        # Verify all outputs exist and have content
        for out_path in outputs:
            assert os.path.exists(out_path), f"Output missing: {out_path}"
            doc = DocxDocument(out_path)
            assert len(doc.paragraphs) > 0, f"Empty document: {out_path}"
            os.unlink(out_path)

        print(f"  ✓ Concurrent batch (5 docs, 3 workers): {total_time:.3f}s total")
        os.unlink(template_path)

    def test_many_variables_performance(self):
        """Template with many (100) variables should still be fast."""
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc = DocxDocument()

        # Create a paragraph with 100 variables
        p = doc.add_paragraph()
        var_text = " ".join(f"{{{v}}}" for v in [f"var{i}" for i in range(100)])
        # Use double-brace format for template variables
        var_text = var_text.replace("{", "{{").replace("}", "}}")
        p.add_run(var_text)
        doc.save(tmp.name)
        template_path = tmp.name

        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        variables = {f"var{i}": f"值{i}" for i in range(100)}

        start = time.time()
        generate_document(template_path, variables, output_path)
        elapsed = time.time() - start

        assert elapsed < 10.0, (
            f"100-variable generation took {elapsed:.2f}s, expected < 10s"
        )
        print(f"  ✓ 100 variables: {elapsed:.3f}s")

        os.unlink(template_path)
        os.unlink(output_path)

    def test_generation_consistency(self):
        """Same input should produce identical output consistently."""
        template_path = _create_template()
        variables = {f"var{i}": f"固定值{i}" for i in range(10)}

        outputs = []
        for _ in range(3):
            out = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            outputs.append(out.name)
            out.close()

        for out_path in outputs:
            generate_document(template_path, variables, out_path)

        # All should have same content
        ref_doc = DocxDocument(outputs[0])
        ref_text = [p.text for p in ref_doc.paragraphs]

        for out_path in outputs[1:]:
            doc = DocxDocument(out_path)
            current_text = [p.text for p in doc.paragraphs]
            assert current_text == ref_text, "Generated content should be consistent"

        for p in outputs:
            os.unlink(p)
        os.unlink(template_path)
