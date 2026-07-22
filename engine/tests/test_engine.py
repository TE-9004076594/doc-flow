"""Tests for the document engine core components."""

import os
import tempfile
import pytest
from docx import Document as DocxDocument

from app.core.template_parser import parse_template, validate_template_structure, VARIABLE_PATTERN
from app.core.renderer import generate_document, detect_unresolved_placeholders
from app.core.logic_blocks import evaluate_condition
from app.core.html_renderer import render_to_html


def _create_test_docx(text: str) -> str:
    """Create a temporary .docx file with the given text."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(tmp.name)
    return tmp.name


class TestTemplateParser:
    """Test template parsing and variable extraction."""

    def test_extract_simple_variables(self):
        text = "公司名称：{{company_name}}\n日期：{{sign_date}}"
        path = _create_test_docx(text)
        result = parse_template(path)
        names = sorted(v.name for v in result.variables)
        assert names == ["company_name", "sign_date"]
        os.unlink(path)

    def test_no_variables(self):
        text = "这是一份普通的文档，没有变量。"
        path = _create_test_docx(text)
        result = parse_template(path)
        assert len(result.variables) == 0
        os.unlink(path)

    def test_detect_condition_blocks(self):
        text = "{{#if:include_terms}}附加条款内容{{/if}}"
        path = _create_test_docx(text)
        result = parse_template(path)
        assert result.has_condition_blocks
        assert result.block_count == 1
        os.unlink(path)

    def test_detect_loop_blocks(self):
        text = "{{#each:items}}产品：{{name}}{{/each}}"
        path = _create_test_docx(text)
        result = parse_template(path)
        assert result.has_loop_blocks
        assert result.block_count == 1
        os.unlink(path)

    def test_validate_balanced_blocks(self):
        text = "{{#if:test}}内容{{/if}}"
        path = _create_test_docx(text)
        errors = validate_template_structure(path)
        assert len(errors) == 0
        os.unlink(path)

    def test_validate_unbalanced_blocks(self):
        text = "{{#if:test}}内容"
        path = _create_test_docx(text)
        errors = validate_template_structure(path)
        assert len(errors) > 0
        os.unlink(path)


class TestLogicBlocks:
    """Test condition and loop block evaluation."""

    def test_truthy_condition(self):
        assert evaluate_condition("include_terms", {"include_terms": True})

    def test_falsy_condition(self):
        assert not evaluate_condition("include_terms", {"include_terms": False})

    def test_equality_condition(self):
        assert evaluate_condition('region == "CN"', {"region": "CN"})
        assert not evaluate_condition('region == "CN"', {"region": "US"})

    def test_numeric_comparison(self):
        assert evaluate_condition("amount > 100", {"amount": 200})
        assert not evaluate_condition("amount > 100", {"amount": 50})

    def test_nested_variable_resolution(self):
        variables = {"items": [{"name": "产品A", "price": 100}]}
        assert evaluate_condition('items.0.name == "产品A"', variables)
        assert not evaluate_condition('items.0.price > 200', variables)


class TestRenderer:
    """Test document generation and variable replacement."""

    def test_basic_variable_replacement(self):
        text = "公司：{{company_name}}"
        template_path = _create_test_docx(text)
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {"company_name": "Acme Corp"}, output_path)

        doc = DocxDocument(output_path)
        assert "Acme Corp" in doc.paragraphs[0].text
        assert "{{company_name}}" not in doc.paragraphs[0].text

        os.unlink(template_path)
        os.unlink(output_path)

    def test_detect_unresolved_placeholders(self):
        text = "已替换：{{resolved_var}}，未替换：{{unresolved_var}}"
        template_path = _create_test_docx(text)
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {"resolved_var": "已填充"}, output_path)
        unresolved = detect_unresolved_placeholders(output_path)
        assert "unresolved_var" in unresolved
        assert "resolved_var" not in unresolved

        os.unlink(template_path)
        os.unlink(output_path)

    def test_multiple_variables(self):
        text = "{{name}}，您的订单 {{order_id}}，金额 {{amount}} 元"
        template_path = _create_test_docx(text)
        output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = output.name
        output.close()

        generate_document(template_path, {
            "name": "张三",
            "order_id": "ORD-2024-001",
            "amount": "1,500.00",
        }, output_path)

        doc = DocxDocument(output_path)
        content = doc.paragraphs[0].text
        assert "张三" in content
        assert "ORD-2024-001" in content
        assert "1,500.00" in content

        os.unlink(template_path)
        os.unlink(output_path)


class TestHtmlRenderer:
    """Test OOXML to HTML conversion."""

    def test_render_simple_paragraph(self):
        text = "Hello World"
        path = _create_test_docx(text)
        html = render_to_html(path)
        assert "Hello World" in html
        assert "<p>" in html or '<p class=' in html
        os.unlink(path)

    def test_render_empty_document(self):
        path = _create_test_docx("")
        html = render_to_html(path)
        assert html is not None
        os.unlink(path)
