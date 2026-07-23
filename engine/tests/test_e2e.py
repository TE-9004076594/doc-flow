"""End-to-end tests for the full document generation pipeline.

Tests: upload template → format extraction → (LLM generation) → rendering → quality check
"""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.models import FieldSchemaSet, FieldSchema, FieldType
from app.core.extractor import extract_format
from app.core.renderer import generate_document, detect_unresolved_placeholders
from app.core.field_validator import validate_content, generate_schema_from_variables
from app.core.quality_checker import run_quality_check


def _create_complex_template() -> str:
    """Create a .docx with diverse formatting for E2E testing."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.left_margin = Cm(3.17)

    # Custom heading style
    heading = doc.styles.add_style('E2EHeading', 1)  # WD_STYLE_TYPE.PARAGRAPH=1
    heading.font.name = 'Arial'
    heading.font.size = Pt(14)
    heading.font.bold = True
    pf = heading.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.line_spacing = 1.5

    # Content with various placeholders
    doc.add_paragraph("合同编号：{{contract_id}}", style='Normal')
    doc.add_paragraph("甲方名称：{{customer_name}}", style='E2EHeading')
    doc.add_paragraph("合同金额：{{amount}} 元")
    doc.add_paragraph("生效日期：{{effective_date}}")

    # Table with placeholders
    table = doc.add_table(rows=3, cols=3, style='Table Grid')
    for i, header in enumerate(['项目', '数量', '金额']):
        table.cell(0, i).text = header
    table.cell(1, 0).text = '{{item_1}}'
    table.cell(1, 1).text = '{{qty_1}}'
    table.cell(1, 2).text = '{{amount_1}}'

    doc.save(tmp.name)
    return tmp.name


class TestE2EGeneration:
    """End-to-end test: upload → extract → render → quality check."""

    @pytest.fixture(autouse=True)
    def setup(self):
        self.template_path = _create_complex_template()
        self.output_path = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        yield
        Path(self.template_path).unlink(missing_ok=True)
        Path(self.output_path).unlink(missing_ok=True)

    # ─── 8.1 Regression Test Set ──────────────────────────────

    def test_format_extraction_completeness(self):
        """Verify extractor captures all major format categories."""
        spec = extract_format(self.template_path)
        assert spec.page.size is not None          # Page setup
        assert "E2EHeading" in spec.styles          # Custom styles
        assert spec.table.style_name is not None    # Table format
        assert len(spec.constraints) > 0            # Constraints

    # ─── 8.2 Pipeline: extract → render → quality ────────────

    def test_full_generation_pipeline(self):
        """Verify the full pipeline: extract format → render → quality check."""
        # Extract
        spec = extract_format(self.template_path)
        assert spec.template_id == ""

        # Render
        variables = {
            "contract_id": "HT-2026-001",
            "customer_name": "测试科技有限公司",
            "amount": "100000",
            "effective_date": "2026-07-22",
            "item_1": "软件开发服务",
            "qty_1": "1",
            "amount_1": "80000",
        }
        result = generate_document(self.template_path, variables, self.output_path)
        assert result == self.output_path
        assert Path(self.output_path).exists()

        # Placeholder check
        unresolved = detect_unresolved_placeholders(self.output_path)
        assert len(unresolved) == 0

        # Quality check
        quality = run_quality_check(self.template_path, self.output_path)
        assert quality is not None

    # ─── 8.3 Format Consistency ≥ 98% ─────────────────────────

    def test_format_consistency_high(self):
        """Verify format consistency score meets the 98% threshold."""
        variables = {
            "contract_id": "HT-2026-001",
            "customer_name": "测试科技有限公司",
            "amount": "100000",
            "effective_date": "2026-07-22",
            "item_1": "软件开发服务",
            "qty_1": "1",
            "amount_1": "80000",
        }
        generate_document(self.template_path, variables, self.output_path)
        quality = run_quality_check(self.template_path, self.output_path)

        # Score should be high since template rendering preserves formatting
        assert quality.consistency_score >= 98.0, (
            f"格式一致率 {quality.consistency_score}% < 98%"
        )

    # ─── 8.4 Placeholder Replacement = 100% ───────────────────

    def test_placeholder_replacement_rate(self):
        """Verify all placeholders are replaced (100% rate)."""
        variables = {
            "contract_id": "HT-2026-001",
            "customer_name": "科技有限公司",
            "amount": "500000",
            "effective_date": "2026-07-22",
            "item_1": "咨询服务",
            "qty_1": "2",
            "amount_1": "150000",
        }
        generate_document(self.template_path, variables, self.output_path)
        unresolved = detect_unresolved_placeholders(self.output_path)

        total_placeholders = 7  # all variables in template
        unresolved_count = len(unresolved)
        replacement_rate = ((total_placeholders - unresolved_count) / total_placeholders) * 100
        assert replacement_rate == 100.0, (
            f"占位符替换率 {replacement_rate}% != 100% (未替换: {unresolved})"
        )

    # ─── Schema Validation Pipeline ───────────────────────────

    def test_schema_generation_and_validation(self):
        """Verify schema generation from variables + content validation."""
        variables = {
            "contract_id": "HT-2026-001",
            "customer_name": "科技有限公司",
            "amount": "500000",
            "effective_date": "2026-07-22",
        }

        # Generate schema from variable names
        schema_set = generate_schema_from_variables("test-e2e", list(variables.keys()))

        # Validate content against schema
        result = validate_content(variables, schema_set)
        assert result.is_valid
        assert len(result.validated_data) == len(variables)

    def test_validation_rejects_missing_required(self):
        """Verify validation catches missing required fields."""
        schema_set = generate_schema_from_variables("test-e2e", ["name", "amount"])
        result = validate_content({"name": "测试"}, schema_set)
        assert not result.is_valid  # amount is missing

    def test_quality_report_structure(self):
        """Verify quality report has expected structure."""
        variables = {
            "contract_id": "HT-2026-001",
            "customer_name": "测试",
            "amount": "100",
            "effective_date": "2026-07-22",
            "item_1": "A",
            "qty_1": "1",
            "amount_1": "100",
        }
        generate_document(self.template_path, variables, self.output_path)
        quality = run_quality_check(self.template_path, self.output_path)
        report = quality.to_dict()

        assert "passed" in report
        assert "consistency_score" in report
        assert "total_issues" in report
        assert "details" in report
        assert "page_setup" in report["details"]
        assert "styles" in report["details"]
        assert "placeholders" in report["details"]
